"""
resume_parser.py — Extract plain text from PDF, DOCX, or TXT resumes.
"""
import re
from pathlib import Path


def extract_text(path: str) -> str:
    """
    Extract clean text from a resume file.
    Supports: .pdf, .docx, .txt, .md

    Returns the full text content, cleaned up for use in AI prompts.
    Raises FileNotFoundError if the file doesn't exist.
    Raises ValueError for unsupported formats.
    """
    p = Path(path)
    if not p.exists():
        raise FileNotFoundError(f"Resume not found: {path}")

    suffix = p.suffix.lower()

    if suffix == ".pdf":
        text = _extract_pdf(p)
    elif suffix == ".docx":
        text = _extract_docx(p)
    elif suffix in (".txt", ".md"):
        text = p.read_text(encoding="utf-8", errors="ignore")
    else:
        raise ValueError(
            f"Unsupported resume format '{suffix}'. Use .pdf, .docx, or .txt"
        )

    return _clean(text)


def _extract_pdf(path: Path) -> str:
    try:
        from pdfminer.high_level import extract_text as pdfminer_extract
        return pdfminer_extract(str(path))
    except ImportError:
        raise ImportError("Run: pip install pdfminer.six")
    except Exception as e:
        raise RuntimeError(f"Could not read PDF: {e}") from e


def _extract_docx(path: Path) -> str:
    try:
        from docx import Document
        doc = Document(str(path))
        parts = []
        for para in doc.paragraphs:
            if para.text.strip():
                parts.append(para.text)
        # Also extract text from tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        parts.append(cell.text)
        return "\n".join(parts)
    except ImportError:
        raise ImportError("Run: pip install python-docx")
    except Exception as e:
        raise RuntimeError(f"Could not read DOCX: {e}") from e


def _clean(text: str) -> str:
    """Remove excessive whitespace while preserving structure."""
    # Normalize line endings
    text = text.replace("\r\n", "\n").replace("\r", "\n")
    # Collapse 3+ blank lines to 2
    text = re.sub(r"\n{3,}", "\n\n", text)
    # Remove non-printable characters (keep newlines and tabs)
    text = re.sub(r"[^\x09\x0A\x0D\x20-\x7E\u00A0-\uFFFF]", " ", text)
    return text.strip()


def get_resume_summary(full_text: str, max_chars: int = 4000) -> str:
    """
    Return a truncated version of the resume for AI prompts.
    Trims from the end to stay within token budgets.
    """
    if len(full_text) <= max_chars:
        return full_text
    # Try to cut at a sentence boundary
    truncated = full_text[:max_chars]
    last_newline = truncated.rfind("\n")
    if last_newline > max_chars * 0.8:
        truncated = truncated[:last_newline]
    return truncated + "\n[... resume truncated for length ...]"
